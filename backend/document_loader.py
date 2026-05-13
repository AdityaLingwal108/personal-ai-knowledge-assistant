from pathlib import Path
from typing import List


def load_document(filepath: str) -> str:
    path = Path(filepath)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _load_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return _load_docx(filepath)
    elif ext == ".txt":
        return _load_txt(filepath)
    else:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: .pdf, .docx, .doc, .txt")


def _load_pdf(filepath: str) -> str:
    try:
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(filepath)
        try:
            pages: List[str] = []
            for page in pdf:
                textpage = page.get_textpage()
                try:
                    text = textpage.get_text_range()
                finally:
                    textpage.close()
                    page.close()
                if text and text.strip():
                    pages.append(text)
        finally:
            pdf.close()

        if not pages:
            raise ValueError(f"No extractable text found in PDF: {filepath}")
        return "\n".join(pages)
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to parse PDF '{Path(filepath).name}': {e}")


def _load_docx(filepath: str) -> str:
    try:
        from docx import Document

        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError(f"No extractable text found in DOCX: {filepath}")
        return "\n".join(paragraphs)
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to parse DOCX '{Path(filepath).name}': {e}")


def _load_txt(filepath: str) -> str:
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        if not text.strip():
            raise ValueError(f"Text file is empty: {filepath}")
        return text
    except ValueError:
        raise
    except Exception as e:
        raise RuntimeError(f"Failed to read TXT '{Path(filepath).name}': {e}")


def chunk_text(text: str, chunk_size: int = 600, overlap: int = 120) -> List[str]:
    # Semantic chunking: split by paragraphs
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks = []
    current_chunk = ""

    def get_overlap(text: str, overlap_size: int) -> str:
        if len(text) <= overlap_size:
            return text
        tail = text[-overlap_size:]
        space_idx = tail.find(" ")
        return tail[space_idx + 1 :] if space_idx != -1 else tail

    for p in paragraphs:
        # If a single paragraph is too large, split by sentences (roughly by '.')
        if len(p) > chunk_size:
            sentences = [s.strip() + "." for s in p.split(".") if s.strip()]
            for s in sentences:
                if len(current_chunk) + len(s) > chunk_size and current_chunk:
                    chunks.append(current_chunk.strip())
                    current_chunk = get_overlap(current_chunk, overlap) + " " + s
                else:
                    current_chunk += " " + s if current_chunk else s
        else:
            if len(current_chunk) + len(p) > chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = get_overlap(current_chunk, overlap) + "\n\n" + p
            else:
                current_chunk += "\n\n" + p if current_chunk else p

    if current_chunk.strip():
        chunks.append(current_chunk.strip())

    # Filter out tiny leftover chunks
    return [c for c in chunks if len(c) >= 50]
